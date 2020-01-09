import uuid
import shutil

from docx import Document
from os.path import join

if __package__ == 'pptx_autobook':
    from pptx_autobook.pptx_proc import pptx_get_outline, pptx_slide_export
else:
    from pptx_proc import pptx_get_outline, pptx_slide_export


def _docx_content_gen(docx, root, imgdir, slidePrefix, level, imgFmt):

    for node in root:

        # Write title
        title = node['title']
        if title != None:
            docx.add_heading(title, level)

        # Put slides
        for content in node['content']:
            table = docx.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            cell.autofit = False
            run = cell.paragraphs[0].add_run()
            run.add_picture(
                join(imgdir, slidePrefix + str(content) + '.' + imgFmt), width=cell.width)

        # Recursive processing
        _docx_content_gen(docx, node['child'],
                          imgdir, slidePrefix, level + 1, imgFmt)

    return docx


def docx_content_gen(docx, outline, imgdir, slidePrefix, imgFmt):
    return _docx_content_gen(docx, [outline], imgdir, slidePrefix, 1, imgFmt)


def docx_autobook(pptPath, inDocPath, outDocPath):

    slidePrefix = '投影片'
    imgFmt = 'PNG'

    # Open/Create document
    doc = Document(inDocPath)

    # For all pptx in sequence
    for p in pptPath:

        # Show processing information
        print('Processing {} ...'.format(p))

        # Get outline from pptx
        outline = pptx_get_outline(p)

        # Export slide images
        uuidStr = str(uuid.uuid4())
        pptx_slide_export(p, uuidStr, imgFmt)

        # Generate document
        doc = docx_content_gen(doc, outline, uuidStr, slidePrefix, imgFmt)

        # Clean up
        shutil.rmtree(uuidStr)

    # Save document
    doc.save(outDocPath)


# Test
if __name__ == '__main__':

    fPath = ['./pptx_file/test.pptx']

    docx_autobook(fPath, None, './docx_file/test_out.docx')
